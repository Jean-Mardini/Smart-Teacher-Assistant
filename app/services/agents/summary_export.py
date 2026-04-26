"""Export structured summary results to DOCX or PDF."""

from __future__ import annotations

import textwrap
from io import BytesIO
from typing import Any, Iterable

import fitz
from docx import Document


def _safe_filename_base(source_documents: list[str]) -> str:
    raw = (source_documents[0] if source_documents else "summary").strip() or "summary"
    safe = "".join(ch if ch.isalnum() else "_" for ch in raw).strip("_")
    return safe or "summary"


def _blocks_from_payload(data: dict[str, Any]) -> list[tuple[str, str]]:
    """Ordered (kind, text) where kind is 'title' | 'h' | 'p' | 'li'."""
    out: list[tuple[str, str]] = []
    out.append(("title", "Summary"))

    if data.get("source_documents"):
        src = data["source_documents"]
        if isinstance(src, list) and src:
            out.append(("p", "Sources: " + " · ".join(str(s) for s in src)))
            if isinstance(data.get("total_pages"), int) and data["total_pages"] > 0:
                out[-1] = ("p", out[-1][1] + f" (~{data['total_pages']} pages declared)")

    summary = (data.get("summary") or "").strip()
    if summary:
        out.append(("h", "Overview"))
        out.append(("p", summary))

    key_points = data.get("key_points") or []
    if isinstance(key_points, list) and key_points:
        out.append(("h", "Key points"))
        for k in key_points:
            out.append(("li", str(k)))

    action_items = data.get("action_items") or []
    if isinstance(action_items, list) and action_items:
        out.append(("h", "Action items"))
        for k in action_items:
            out.append(("li", str(k)))

    formulas = data.get("formulas") or []
    if isinstance(formulas, list) and formulas:
        out.append(("h", "Formulas"))
        for f in formulas:
            out.append(("li", str(f)))

    glossary = data.get("glossary") or []
    if isinstance(glossary, list) and glossary:
        out.append(("h", "Glossary"))
        for g in glossary:
            if isinstance(g, dict):
                term = str(g.get("term") or "").strip()
                defin = str(g.get("definition") or "").strip()
                if term:
                    out.append(("p", f"{term}: {defin}" if defin else term))

    image_notes = data.get("image_notes") or []
    if isinstance(image_notes, list) and image_notes:
        out.append(("h", "Image notes"))
        for n in image_notes:
            out.append(("li", str(n)))

    processing_notes = data.get("processing_notes") or []
    if isinstance(processing_notes, list) and processing_notes:
        out.append(("h", "Processing notes"))
        out.append(("p", " ".join(str(n) for n in processing_notes)))

    return out


def summary_payload_to_docx_bytes(data: dict[str, Any]) -> tuple[bytes, str]:
    doc = Document()
    blocks = _blocks_from_payload(data)
    for kind, text in blocks:
        if kind == "title":
            doc.add_heading(text, level=0)
        elif kind == "h":
            doc.add_heading(text, level=2)
        elif kind == "p":
            doc.add_paragraph(text)
        elif kind == "li":
            doc.add_paragraph(text, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    base = _safe_filename_base(list(data.get("source_documents") or []))
    return buf.getvalue(), f"{base}.docx"


def _pdf_emit_blocks(blocks: Iterable[tuple[str, str]]) -> bytes:
    doc = fitz.open()
    w_pt, h_pt = fitz.paper_size("a4")
    margin = 56
    max_y = h_pt - margin
    line_w_chars = 92

    page: fitz.Page | None = None
    y = margin

    def new_page() -> fitz.Page:
        nonlocal page, y
        page = doc.new_page(width=w_pt, height=h_pt)
        y = margin
        return page

    def ensure_room(line_h: float) -> fitz.Page:
        nonlocal page, y
        if page is None or y + line_h > max_y:
            return new_page()
        return page

    for kind, raw in blocks:
        text = raw.replace("\r\n", "\n").replace("\r", "\n")
        if kind == "title":
            fs = 16
            lh = fs * 1.4
            p = ensure_room(lh)
            p.insert_text((margin, y + fs), text, fontsize=fs, fontname="hebo")
            y += lh + 6
        elif kind == "h":
            fs = 12
            lh = fs * 1.35
            p = ensure_room(lh + 4)
            p.insert_text((margin, y + fs), text, fontsize=fs, fontname="hebo")
            y += lh + 8
        elif kind in ("p", "li"):
            fs = 10
            lh = fs * 1.38
            bullet = "• "
            for i, para in enumerate(text.split("\n")):
                lead = ""
                if kind == "li":
                    lead = bullet if i == 0 else "  "
                chunk = (lead + para).strip() if para.strip() else (lead if kind == "li" else "")
                lines = textwrap.wrap(
                    chunk,
                    width=line_w_chars,
                    break_long_words=True,
                    replace_whitespace=False,
                ) or ([chunk] if chunk else [lead if kind == "li" else " "])
                for line in lines:
                    p = ensure_room(lh)
                    p.insert_text((margin, y + fs), line, fontsize=fs, fontname="helv")
                    y += lh
            y += 4

    if doc.page_count == 0:
        new_page()
        page.insert_text((margin, margin + 10), "(Empty summary)", fontsize=10, fontname="helv")

    out = doc.tobytes()
    doc.close()
    return out


def summary_payload_to_pdf_bytes(data: dict[str, Any]) -> tuple[bytes, str]:
    blocks = _blocks_from_payload(data)
    body = _pdf_emit_blocks(blocks)
    base = _safe_filename_base(list(data.get("source_documents") or []))
    return body, f"{base}.pdf"
