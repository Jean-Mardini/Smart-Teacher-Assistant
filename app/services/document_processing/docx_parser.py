from docx import Document
from .cleaners import clean_text
from .tables import extract_docx_tables
import os


# Style names that python-docx uses for headings
_HEADING_STYLES = {
    'heading 1', 'heading 2', 'heading 3',
    'title', 'subtitle',
    # French variants
    'titre', 'sous-titre',
}

_APPROX_CHARS_PER_PAGE = 3000


def extract_docx_text(filepath: str, document_id: str = "doc"):
    """
    Extract text, tables, and images from a DOCX file.

    Returns:
        pages:     list of {"page": N, "text": "..."} — approximate page groupings
        full_text: complete cleaned document text
        tables:    list of extracted table dicts
        images:    list of image dicts with path (no binary)
    """

    doc = Document(filepath)

    full_text = ""
    page_number = 1
    page_buffer = ""
    pages = []

    for para in doc.paragraphs:
        # python-docx: paragraphs in some templates have no assigned style (style is None).
        style_obj = getattr(para, "style", None)
        style_name = ((getattr(style_obj, "name", None) or "") or "").lower()
        raw = para.text.strip()

        if not raw:
            continue

        # Uppercase Heading 1 / Title so structure_extraction's ALL CAPS rule fires
        if 'heading 1' in style_name or style_name == 'title':
            line = raw.upper()
        else:
            line = raw

        cleaned = clean_text(line)
        full_text += cleaned + "\n"
        page_buffer += cleaned + "\n"

        if len(page_buffer) >= _APPROX_CHARS_PER_PAGE:
            pages.append({"page": page_number, "text": page_buffer.strip()})
            page_buffer = ""
            page_number += 1

    if page_buffer.strip():
        pages.append({"page": page_number, "text": page_buffer.strip()})

    tables = extract_docx_tables(doc)

    # Images
    images = []
    image_counter = 0
    img_dir = os.path.join("outputs", "images", document_id)

    for shape in doc.inline_shapes:
        try:
            pic = shape._inline.graphic.graphicData.pic
            r_id = pic.blipFill.blip.embed
            image_part = doc.part.related_parts[r_id]

            image_counter += 1
            img_id = f"img_{image_counter}"
            os.makedirs(img_dir, exist_ok=True)
            img_path = os.path.join(img_dir, f"{img_id}.png")
            with open(img_path, "wb") as f:
                f.write(image_part.blob)

            # Alt text: prefer descr, fall back to name
            doc_pr = shape._inline.docPr
            caption = doc_pr.get("descr") or doc_pr.get("name") or ""

            images.append({
                "image_id": img_id,
                "page": page_number,
                "caption": caption,
                "path": img_path,
            })
        except Exception:
            continue

    return pages, full_text, tables, images
